"""DOCX parser (python-docx). Joins non-empty paragraphs with newlines."""
from __future__ import annotations

import io

from docx import Document


def parse(data: bytes):
    from . import ParsedDocument

    document = Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(paragraphs).strip()
    return ParsedDocument(text=text, metadata={"paragraph_count": len(paragraphs)})
