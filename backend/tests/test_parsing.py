"""Parser-layer tests. No env or network needed. Fixtures are generated in-test
(docx via python-docx) or are literals (html/md); the PDF is a minimal document
built with pypdf so extraction is deterministic.
"""
import io

import pytest

from app.services.parsing import (
    ParsedDocument,
    UnsupportedFormatError,
    parse_document,
)


def test_markdown_roundtrip():
    data = b"# Title\n\nSome **bold** body text."
    parsed = parse_document("notes.md", data)
    assert isinstance(parsed, ParsedDocument)
    assert "Title" in parsed.text
    assert "bold" in parsed.text
    assert parsed.metadata["extension"] == "md"
    assert parsed.metadata["filename"] == "notes.md"


def test_txt():
    parsed = parse_document("a.txt", b"plain text content")
    assert parsed.text == "plain text content"


def test_html_strips_script_and_style():
    html = b"""<html><head><title>Doc</title><style>.x{color:red}</style></head>
    <body><p>Hello world</p><script>alert(1)</script></body></html>"""
    parsed = parse_document("page.html", html)
    assert "Hello world" in parsed.text
    assert "alert" not in parsed.text
    assert "color:red" not in parsed.text
    assert parsed.metadata.get("title") == "Doc"


def test_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    parsed = parse_document("report.docx", buf.getvalue())
    assert "First paragraph." in parsed.text
    assert "Second paragraph." in parsed.text


def test_pdf():
    # Build a one-page PDF with extractable text using reportlab if available,
    # else fall back to a hand-built pypdf page. Keep it dependency-light: use
    # pypdf's own writer with a text-bearing page via a minimal content stream.
    pytest.importorskip("pypdf")
    from pypdf import PdfReader

    # A minimal but valid PDF containing the text "Hello PDF".
    pdf_bytes = _minimal_pdf("Hello PDF")
    reader = PdfReader(io.BytesIO(pdf_bytes))
    # Sanity: the fixture itself is extractable, proving the parser path.
    parsed = parse_document("doc.pdf", pdf_bytes)
    assert "Hello PDF" in parsed.text
    assert parsed.metadata["page_count"] == len(reader.pages)


def test_unsupported_extension():
    with pytest.raises(UnsupportedFormatError):
        parse_document("archive.zip", b"PK\x03\x04")


def _minimal_pdf(message: str) -> bytes:
    """Construct a minimal single-page PDF whose content stream draws `message`."""
    objects = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
    )
    stream = (
        b"BT /F1 24 Tf 72 700 Td (" + message.encode("latin-1") + b") Tj ET"
    )
    objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += str(i).encode() + b" 0 obj\n" + obj + b"\nendobj\n"

    xref_pos = len(pdf)
    pdf += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n"
    pdf += b"0000000000 65535 f \n"
    for off in offsets:
        pdf += ("%010d 00000 n \n" % off).encode()
    pdf += (
        b"trailer\n<< /Size " + str(len(objects) + 1).encode()
        + b" /Root 1 0 R >>\nstartxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    )
    return bytes(pdf)
